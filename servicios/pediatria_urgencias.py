import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
from datetime import date
from datetime import datetime
from zoneinfo import ZoneInfo
import unicodedata
import json
import re
from io import BytesIO
from html import escape
from pathlib import Path
from difflib import SequenceMatcher
from deep_translator import GoogleTranslator
from pypdf import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from core.calculos import calcular_edad, edad_en_meses
from core.clasificacion import grupo_etario
from core.texto import sexo_texto
from herramientas.neurodesarrollo import obtener_neurodesarrollo
from herramientas.antropometria import calcular_imc
from herramientas.superficie_corporal import calcular_sc
from herramientas.oms_full import (
    zscore_peso_edad,
    zscore_talla_edad,
    zscore_imc_edad,
    zscore_peso_talla,
    zscore_pc_edad
)

# 🔥 IMPORTANTE (NUEVO)
from herramientas.diagnostico_nutricional import (
    diagnostico_menor_5,
    diagnostico_mayor_5
)


BASE_DIR = Path(__file__).resolve().parent.parent
HISTORIAS_PATH = BASE_DIR / "data" / "historias_pediatria_urgencias.jsonl"
PLANES_PATOLOGIA_PATH = BASE_DIR / "data" / "planes_manejo_pediatria_urgencias.json"
DOSIS_MEDICACION_PATH = BASE_DIR / "data" / "dosis_medicacion_pediatria_urgencias.json"
BOGOTA_TZ = ZoneInfo("America/Bogota")

ANTECEDENTES_DEFAULT = """NEONATALES: PRODUCTO XX GESTACIÓN, MADRE DE XX AÑOS, CONTROLADO, SIN COMPLICACIONES, STORCH NEGATIVO Y ECOGRAFÍAS ANTENATALES: NACE VÍA VAGINAL/ CESAREA A LAS XX SEM A TÉRMINO. EGRESO CONJUNTO, PESO XXXX GR - TALLA XX CM.
INMUNOLÓGICOS: VACUNAS AL DÍA SEGÚN PAI.
ALIMENTACIÓN: ACORDE A EDAD.
PATOLÓGICOS: NIEGA.
HOSPITALARIOS: NIEGA.
FARMACOLÓGICOS: NIEGA.
TRAUMÁTICOS: NIEGA.
TOXICOLÓGICO: NIEGA EXPOSICIÓN A HUMO DE LEÑA O CIGARRILLO.
ALÉRGICOS: NIEGA.
TRANSFUSIONALES: NIEGA.
QUIRÚRGICOS: NIEGA.
FAMILIARES: PADRE SANO, MADRE SANA, NO CONSANGUÍNEOS.
HEMOCLASIFICACIÓN: O POSITIVO.
PSICOSOCIALES: VIVIENDA CON TODOS LOS SERVICIOS."""

EXAMEN_DEFAULT = """PACIENTE LUCE EN BUEN ESTADO GENERAL, ALERTA, BUEN ESTADO DE HIDRATACIÓN, AFEBRIL.

CABEZA: NORMOCÉFALA, SIN LESIONES.
OJOS: CONJUNTIVAS ROSADAS, ESCLERAS ANICTÉRICAS.
OÍDOS: SIN ALTERACIONES.
NARIZ: PERMEABLE.
OROFARINGE: MUCOSAS HÚMEDAS, SIN LESIONES.
CUELLO: MÓVIL, SIN ADENOPATÍAS.
TÓRAX: SIMÉTRICO, NORMOEXPANSIBLE, SIN TIRAJES.
CARDIOPULMONAR: RUIDOS CARDIACOS RÍTMICOS, SIN SOPLOS, SIN AGREGADOS PULMONARES, OXIMETRIAS ADECUADAS AL AIRE AMBIENTE.
ABDOMEN: NO DISTENDIDO, RSHS PRESENTES, BLANDO, NO DOLOROSO, SIN SIGNOS DE IRRITACIÓN PERITONEAL.
GENITALES: INFANTILES ACORDES NORMOCONFIGURADOS. 
EXTREMIDADES: EUTROFICAS, SIN EDEMAS.
NEUROLÓGICO: ALERTA, NO FOCALIZACIONES, SIN SIGNOS DE MENINGISMO, ROT NORMALES, FUERZA MUSCULAR CONSERVADA, SIN DÉFICIT.
PIEL: ROSADA, BIEN PERFUNDIDA, SIN LESIONES.
"""

PLAN_DEFAULT = """- HOSPITALIZACION PEDIATRICA
- DIETA NORMAL ACORDE A LA EDAD
- CATETER SELLADO
- LACTATO DE RINGER A 45 CC/HORA IV
- ACETAMINOFEN: 150 MG VO CADA 6 HORAS (SI FIEBRE O DOLOR)
- DIPIRONA: 500 MG CADA 8 HORAS IV (SI FIEBRE O DOLOR NO CEDE CON ACETAMINOFEN)
- SS PARACLINICOS DE EXTENSIÓN
- CONTROL DE LÍQUIDOS ADMINISTRADOS - ELIMINADOS
- CONTROL DE SIGNOS VITALES, AVISAR CAMBIOS."""

REVISION_DEFAULT = "NIEGA OTROS SINTOMAS/SIGNOS A LOS YA MENCIONADOS."

DOSIS_MEDICACION_DEFAULTS = {
    "ACETAMINOFEN": {
        "mg_kg_dosis": 15.0,
        "intervalo_horas": 6,
        "max_mg": 650.0,
        "via": "VO",
        "indicacion": "SI FIEBRE O DOLOR",
    },
    "DIPIRONA": {
        "mg_kg_dosis": 20.0,
        "intervalo_horas": 8,
        "max_mg": 1000.0,
        "via": "IV",
        "indicacion": "SI FIEBRE O DOLOR NO CEDE CON ACETAMINOFEN",
    },
}

PLANES_PATOLOGIA_LABELS = {
    "J21": "BRONQUIOLITIS (J21)",
    "J18": "NEUMONÍA (J18)",
    "J05_J04": "LARINGITIS / CRUP (J05 / J04)",
    "H66": "OTITIS MEDIA AGUDA (H66)",
    "J00_J06_J03": "IRA ALTA / NASOFARINGITIS / AMIGDALITIS (J00 / J06 / J03)",
    "A09": "GASTROENTERITIS (A09)",
}

PLANES_PATOLOGIA_DEFAULTS = {
    "J21": """- LAVADOS NASALES FRECUENTES Y ASPIRACION DE SECRECIONES SEGUN NECESIDAD.
- OXIGENO SUPLEMENTARIO SI SATURACION < 92% O SIGNOS DE DIFICULTAD RESPIRATORIA.
- MONITORIZACION DE SIGNOS VITALES Y TRABAJO RESPIRATORIO.
- CONTROL DE INGESTA Y ELIMINACION.""",
    "J18": """- OXIGENO SUPLEMENTARIO SEGUN REQUERIMIENTO Y META DE SATURACION > 92%.
- CEFTRIAXONA 50 MG/KG/DIA IV CADA 24 HORAS.
- CONTROL DE SIGNOS VITALES Y DIFICULTAD RESPIRATORIA.
- TOMAR PARACLINICOS DE CONTROL SEGUN EVOLUCION.""",
    "J05_J04": """- OBSERVACIÓN PEDIATRICA
- DIETA NORMAL ACORDE A LA EDAD (FRACCIONADA Y CON PRECAUCIÓN)
- CATETER SELLADO
{LINEA_ACETAMINOFEN}
{LINEA_DIPIRONA}
{LINEA_DEXAMETASONA_LARINGITIS}
- ADRENALINA NEBULIZADA CICLO X 3
- SS PARACLINICOS DE EXTENSIÓN
- VIGILANCIA DE ESTRIDOR, TIRAJES Y SATURACION
- CONTROL DE SIGNOS VITALES, AVISAR CAMBIOS
- REVALORACIÓN""",
    "H66": """- MANEJO ANALGESICO Y VIGILANCIA CLINICA.
- AMOXICILINA 90 MG/KG/DIA VO DIVIDIDA CADA 12 HORAS.""",
    "J00_J06_J03": """- MANEJO SINTOMATICO Y VIGILANCIA DE SIGNOS DE ALARMA.
- LAVADOS NASALES Y ADECUADA HIDRATACION ORAL.""",
    "A09": """- HIDRATACION ORAL CON SRO SEGUN TOLERANCIA Y ESTADO CLINICO.
- VIGILAR SIGNOS DE DESHIDRATACION Y GASTO URINARIO.
- ONDANSETRON SI VOMITO SEGUN PESO Y CRITERIO CLINICO.""",
}

FORM_DEFAULTS = {
    "nombre_1": "",
    "tipo_documento_1": None,
    "documento_1": "",
    "fecha_1": None,
    "sexo_1": None,
    "eps_1": "",
    "informante_1": "",
    "proveniente_1": "",
    "motivo_1": "",
    "enfermedad_1": "",
    "antecedentes_1": ANTECEDENTES_DEFAULT,
    "neuro_1": "",
    "neuro_prev": None,
    "revision": REVISION_DEFAULT,
    "ta": "",
    "fc": 0.0,
    "fr": 0.0,
    "sat": 0.0,
    "temp": 0.0,
    "peso": 0.0,
    "talla": 0.0,
    "pc": 0.0,
    "scq_pct": 0.0,
    "examen": EXAMEN_DEFAULT,
    "analisis": "",
    "analisis_base": "",
    "busqueda_cie10": "",
    "dx_cie10": None,
    "obs_dx": "",
    "plan": PLAN_DEFAULT,
    "plan_base": PLAN_DEFAULT,
    "paraclinicos_texto": "",
    "paraclinicos_auto": "",
    "paraclinicos_pdf_sig": "",
    "imagenes_texto": "",
    "imagenes_auto": "",
    "imagenes_pdf_sig": "",
}

EQUIVALENCIAS_BUSQUEDA = {
    "asma": "asthma",
    "agud": "acute",
    "aguda": "acute",
    "agudo": "acute",
    "infec": "infection",
    "infec": "infection",
    "infect": "infection",
    "vias": "tract",
    "via": "tract",
    "resp": "respiratory",
    "bonquioli": "bronchiolitis",
    "bronquioli": "bronchiolitis",
    "bronquiolit": "bronchiolitis",
    "bronquiol": "bronchiolitis",
    "bronquit": "bronchitis",
    "nasofarin": "nasopharyngitis",
    "nasofaring": "nasopharyngitis",
    "nasofar": "nasopharyngitis",
    "rinofarin": "rhinopharyngitis",
    "rinofaring": "rhinopharyngitis",
    "faring": "pharyngitis",
    "faringo": "pharyngitis",
    "amigdal": "tonsillitis",
    "adenoid": "adenoid",
    "laring": "laryngitis",
    "traque": "tracheitis",
    "neumon": "pneumonia",
    "otit": "otitis",
    "sinus": "sinusitis",
    "gastroenter": "gastroenteritis",
    "diarre": "diarrhoea",
    "vomit": "vomiting",
    "apendic": "appendicitis",
    "celulit": "cellulitis",
    "dermatit": "dermatitis",
    "urin": "urinary",
    "cistit": "cystitis",
    "pielon": "pyelonephritis",
    "convuls": "convulsion",
    "cuerpo": "body",
    "extran": "foreign",
    "extra": "foreign",
    "esof": "oesoph",
    "esofag": "oesophag",
    "esofago": "oesophagus",
    "esofagu": "oesophag",
    "estomac": "stomach",
    "intestin": "intestin",
    "gargant": "throat",
    "oido": "ear",
    "pulmon": "lung",
}

ALIAS_DESCRIPCION = {
    "acute": "agudo aguda agud",
    "chronic": "cronico cronica cron",
    "infection": "infeccion infec infect",
    "infections": "infeccion infecciones infec infect",
    "infective": "infeccioso infecciosa infect",
    "inflammatory": "inflamatorio inflamatoria inflama",
    "respiratory": "respiratorio respiratoria resp",
    "upper": "superior alta altas sup",
    "lower": "inferior inferiores baja bajas inf",
    "tract": "via vias tracto",
    "multiple": "multiple multiples",
    "unspecified": "no especificado no especificada inespecifico inespecifica",
    "body": "cuerpo",
    "foreign": "extrano extra",
    "oesophagus": "esofago esofag esof",
    "oesophageal": "esofagico esofag",
    "esophagus": "esofago esofag esof",
    "nasopharyngitis": "nasofaringitis nasofarin nasofaring nasofar",
    "nasopharyngeal": "nasofaringea nasofaringeo nasofar",
    "pharyngitis": "faringitis faringo faring",
    "tonsillitis": "amigdalitis amigdal",
    "laryngitis": "laringitis laring",
    "tracheitis": "traqueitis traque",
    "bronchiolitis": "bronquiolitis bonquioli bronquioli bronquiolit bronquiol",
    "bronchitis": "bronquitis bronquit",
    "pneumonia": "neumonia neumon",
    "otitis": "otitis otit oido",
    "sinusitis": "sinusitis sinus",
    "gastroenteritis": "gastroenteritis gastroenter",
    "appendicitis": "apendicitis apendic",
    "diarrhoea": "diarrea diarre",
    "urinary": "urinario urinaria urin",
    "cystitis": "cistitis cistit",
    "pyelonephritis": "pielonefritis pielon",
    "convulsion": "convulsion convuls",
}

MESES_ESP = {
    "ENE": "01",
    "FEB": "02",
    "MAR": "03",
    "ABR": "04",
    "MAY": "05",
    "JUN": "06",
    "JUL": "07",
    "AGO": "08",
    "SEP": "09",
    "SEPT": "09",
    "OCT": "10",
    "NOV": "11",
    "DIC": "12",
}


@st.cache_data
def cargar_cie10():
    df = pd.read_csv("data/cie10_who.csv")
    df["label_normalized"] = df["label"].map(normalizar_texto)
    df["description_normalized"] = df["description"].astype(str).map(normalizar_texto)
    df["code_normalized"] = df["code"].astype(str).map(normalizar_texto)
    df["search_text"] = df.apply(
        lambda row: construir_texto_busqueda(
            row["code_normalized"],
            row["description_normalized"],
            row["label_normalized"]
        ),
        axis=1
    )
    df["search_tokens"] = df["search_text"].map(tokenizar_texto)
    return df

def normalizar_texto(texto):
    texto = "" if texto is None else str(texto)
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    return texto.lower().strip()


def tokenizar_texto(texto):
    return [t for t in re.split(r"[^a-z0-9]+", normalizar_texto(texto)) if t]


def construir_texto_busqueda(code, description, label):
    partes = [code, description, label]
    for termino, alias in ALIAS_DESCRIPCION.items():
        if termino in description:
            partes.append(alias)
    return " ".join(partes)


@st.cache_data(show_spinner=False)
def traducir_cie10_descripcion(descripcion):
    try:
        return GoogleTranslator(source="en", target="es").translate(str(descripcion))
    except Exception:
        return str(descripcion)


@st.cache_data(show_spinner=False)
def traducir_texto_a_ingles(texto):
    try:
        return GoogleTranslator(source="es", target="en").translate(str(texto))
    except Exception:
        return str(texto)


def expandir_terminos_busqueda(texto):
    termino = normalizar_texto(texto)
    terminos = {termino} if termino else set()
    encontro_equivalencia = False

    for origen, destino in sorted(EQUIVALENCIAS_BUSQUEDA.items(), key=lambda item: len(item[0]), reverse=True):
        if origen in termino:
            encontro_equivalencia = True
            terminos.add(destino)
            break

    if len(termino) >= 5 and not encontro_equivalencia:
        traduccion = normalizar_texto(traducir_texto_a_ingles(texto))
        if traduccion:
            terminos.add(traduccion)

    return {t for t in terminos if t}


def construir_grupos_busqueda(texto):
    texto_normalizado = normalizar_texto(texto)
    if "%" in texto_normalizado:
        partes = [p for p in texto_normalizado.split("%") if normalizar_texto(p)]
    else:
        partes = [p for p in texto_normalizado.split() if normalizar_texto(p)]

    if len(partes) <= 1:
        return []

    if not partes:
        return []

    grupos = []
    tiene_vias = any(parte.startswith("via") or parte.startswith("vias") for parte in partes)

    for parte in partes:
        opciones = set(expandir_terminos_busqueda(parte))
        if tiene_vias and (parte == "via" or parte == "vias"):
            opciones.update({"tract", "respiratory", "airway", "airways"})
        if tiene_vias and parte == "inf":
            opciones = {"inferior", "inferiores", "lower", "baja", "bajas"}
        if tiene_vias and parte == "sup":
            opciones = {"superior", "superiores", "upper", "alta", "altas"}
        grupos.append(list(opciones))

    return grupos


def aplanar_grupos_busqueda(grupos):
    terminos = []
    for grupo in grupos:
        for termino in grupo:
            if termino and termino not in terminos:
                terminos.append(termino)
    return terminos


def coincide_termino(term, row):
    if not term:
        return False

    code = row["code_normalized"]
    search_text = row["search_text"]
    search_tokens = row["search_tokens"]

    if code.startswith(term):
        return True

    if len(term) <= 3:
        return any(token.startswith(term) for token in search_tokens)

    return term in search_text or any(token.startswith(term) for token in search_tokens)


def coincide_grupos(row, grupos):
    if not grupos:
        return True

    return all(
        any(coincide_termino(opcion, row) for opcion in opciones_fragmento)
        for opciones_fragmento in grupos
        if opciones_fragmento
    )


def puntuar_diagnostico(row, terminos, grupos=None):
    score = 0
    code = row["code_normalized"]
    desc = row["description_normalized"]
    label = row["label_normalized"]
    search_tokens = row["search_tokens"]

    for termino in terminos:
        if code.startswith(termino):
            score = max(score, 100)
        elif desc.startswith(termino):
            score = max(score, 95)
        elif f" {termino}" in desc or desc == termino:
            score = max(score, 92)
        elif any(token.startswith(termino) for token in search_tokens):
            score = max(score, 90)
        elif termino in desc:
            score = max(score, 85)
        elif termino in label:
            score = max(score, 75)
        else:
            similitud = max(
                SequenceMatcher(None, termino, code).ratio(),
                SequenceMatcher(None, termino, desc[: max(len(termino) + 8, 12)]).ratio(),
            )
            if similitud >= 0.72:
                score = max(score, int(similitud * 100))

    if grupos:
        if coincide_grupos(row, grupos):
            score = max(score, 120 + len(grupos))

    return score


def limpiar_formulario():
    for key, value in FORM_DEFAULTS.items():
        st.session_state[key] = value
    # Los file_uploader no se pueden resetear asignándoles valores directos.
    st.session_state.pop("pdf_paraclinicos_uploader", None)
    st.session_state.pop("pdf_imagenes_uploader", None)


def solicitar_limpieza_formulario():
    st.session_state["_limpiar_formulario"] = True


def guardar_historia(datos):
    HISTORIAS_PATH.parent.mkdir(parents=True, exist_ok=True)
    with HISTORIAS_PATH.open("a", encoding="utf-8") as f:
        f.write(json.dumps(datos, ensure_ascii=False) + "\n")


def cargar_planes_patologia():
    planes = PLANES_PATOLOGIA_DEFAULTS.copy()
    if PLANES_PATOLOGIA_PATH.exists():
        try:
            with PLANES_PATOLOGIA_PATH.open("r", encoding="utf-8") as f:
                datos = json.load(f)
            if isinstance(datos, dict):
                for clave, valor in datos.items():
                    if clave == "__labels__":
                        continue
                    if isinstance(valor, str):
                        planes[clave] = valor
        except Exception:
            pass
    return planes


def cargar_labels_planes_patologia():
    labels = PLANES_PATOLOGIA_LABELS.copy()
    planes = cargar_planes_patologia()
    for clave in planes.keys():
        labels.setdefault(clave, clave.replace("_", " "))

    if PLANES_PATOLOGIA_PATH.exists():
        try:
            with PLANES_PATOLOGIA_PATH.open("r", encoding="utf-8") as f:
                datos = json.load(f)
            labels_guardados = datos.get("__labels__", {}) if isinstance(datos, dict) else {}
            if isinstance(labels_guardados, dict):
                for clave, valor in labels_guardados.items():
                    if isinstance(valor, str) and valor.strip():
                        labels[clave] = valor
        except Exception:
            pass
    return labels


def guardar_planes_patologia(planes, labels=None):
    PLANES_PATOLOGIA_PATH.parent.mkdir(parents=True, exist_ok=True)
    payload = dict(planes)
    if labels:
        payload["__labels__"] = labels
    with PLANES_PATOLOGIA_PATH.open("w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def cargar_dosis_medicacion():
    dosis = json.loads(json.dumps(DOSIS_MEDICACION_DEFAULTS))
    if DOSIS_MEDICACION_PATH.exists():
        try:
            with DOSIS_MEDICACION_PATH.open("r", encoding="utf-8") as f:
                datos = json.load(f)
            if isinstance(datos, dict):
                for medicamento, configuracion in datos.items():
                    if isinstance(configuracion, dict):
                        if medicamento not in dosis:
                            dosis[medicamento] = {}
                        dosis[medicamento].update(configuracion)
        except Exception:
            pass
    return dosis


def guardar_dosis_medicacion(dosis):
    DOSIS_MEDICACION_PATH.parent.mkdir(parents=True, exist_ok=True)
    with DOSIS_MEDICACION_PATH.open("w", encoding="utf-8") as f:
        json.dump(dosis, f, ensure_ascii=False, indent=2)


def obtener_clave_plan_patologia(codigo):
    if not codigo:
        return ""
    if codigo.startswith("J21"):
        return "J21"
    if codigo.startswith("J18"):
        return "J18"
    if codigo.startswith("J05") or codigo.startswith("J04"):
        return "J05_J04"
    if codigo.startswith("H66"):
        return "H66"
    if codigo.startswith("J00") or codigo.startswith("J06") or codigo.startswith("J03"):
        return "J00_J06_J03"
    if codigo.startswith("A09"):
        return "A09"
    return ""


def es_diagnostico_laringitis(diagnostico):
    texto = normalizar_texto(diagnostico)
    return "laringitis" in texto or "laring" in texto


class SafeFormatDict(dict):
    def __missing__(self, key):
        return "{" + key + "}"


@st.cache_resource
def cargar_ocr():
    try:
        from rapidocr_onnxruntime import RapidOCR
        return RapidOCR()
    except Exception:
        return None


def extraer_texto_ocr_de_imagenes(page):
    ocr = cargar_ocr()
    if ocr is None:
        return ""

    bloques = []
    try:
        imagenes = list(page.images)
    except Exception:
        imagenes = []

    for image_file in imagenes:
        try:
            resultado, _ = ocr(image_file.data)
        except Exception:
            continue

        if not resultado:
            continue

        lineas = sorted(
            resultado,
            key=lambda item: (
                min(p[1] for p in item[0]),
                min(p[0] for p in item[0]),
            )
        )
        textos = [str(item[1]).strip() for item in lineas if str(item[1]).strip()]
        if textos:
            bloques.append("\n".join(textos))

    return "\n\n".join(bloques).strip()


def extraer_texto_pdf(pdf_file):
    if not pdf_file:
        return ""

    try:
        pdf_bytes = pdf_file.getvalue()
        reader = PdfReader(BytesIO(pdf_bytes))
        paginas = []
        tiene_imagenes = False
        for page in reader.pages:
            try:
                if len(page.images) > 0:
                    tiene_imagenes = True
            except Exception:
                pass
            texto = page.extract_text() or ""
            if texto.strip():
                paginas.append(texto)
        if paginas:
            return "\n".join(paginas)
        if tiene_imagenes:
            paginas_ocr = []
            for page in reader.pages:
                texto_ocr = extraer_texto_ocr_de_imagenes(page)
                if texto_ocr:
                    paginas_ocr.append(texto_ocr)
            if paginas_ocr:
                return "\n\n".join(paginas_ocr)
            return "__PDF_ESCANEADO_SIN_TEXTO__"
        return ""
    except Exception:
        return ""


def extraer_fecha_principal(texto):
    texto = normalizar_texto_para_reporte(texto)
    match_ymd = re.search(r"\b(\d{4})[-/](\d{1,2})[-/](\d{1,2})\b", texto)
    if match_ymd:
        return f"{match_ymd.group(3).zfill(2)}/{match_ymd.group(2).zfill(2)}/{match_ymd.group(1)}"

    idx_ingreso = re.search(r"FECHA DE INGRESO", texto, flags=re.IGNORECASE)
    if idx_ingreso:
        subtexto = texto[idx_ingreso.start(): idx_ingreso.start() + 220]
        match_ymd_local = re.search(r"\b(\d{4})[-/](\d{1,2})[-/](\d{1,2})\b", subtexto)
        if match_ymd_local:
            return f"{match_ymd_local.group(3).zfill(2)}/{match_ymd_local.group(2).zfill(2)}/{match_ymd_local.group(1)}"
        match_mes_local = re.search(
            r"\b(\d{1,2})[-/ ]([A-ZÁÉÍÓÚa-záéíóú]{3,5})\.?[-/ ](\d{4})\b",
            subtexto,
            flags=re.IGNORECASE
        )
        if match_mes_local:
            dia = match_mes_local.group(1).zfill(2)
            mes_txt = normalizar_texto(match_mes_local.group(2)).upper()
            anio = match_mes_local.group(3)
            mes = MESES_ESP.get(mes_txt[:4]) or MESES_ESP.get(mes_txt[:3])
            if mes:
                return f"{dia}/{mes}/{anio}"

    match_mes = re.search(
        r"\b(\d{1,2})[-/ ]([A-ZÁÉÍÓÚa-záéíóú]{3,5})\.?[-/ ](\d{4})\b",
        texto,
        flags=re.IGNORECASE
    )
    if match_mes:
        dia = match_mes.group(1).zfill(2)
        mes_txt = normalizar_texto(match_mes.group(2)).upper()
        anio = match_mes.group(3)
        mes = MESES_ESP.get(mes_txt[:4]) or MESES_ESP.get(mes_txt[:3])
        if mes:
            return f"{dia}/{mes}/{anio}"

    patrones_preferidos = [
        r"RESULTADO[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"REALIZACI[ÓO]N[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"FECHA DE INGRESO[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"FECHA VALIDACION[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"FECHA RECEPCI[ÓO]N[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"FECHA RECEPCION[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
    ]
    for patron in patrones_preferidos:
        match = re.search(patron, texto, flags=re.IGNORECASE)
        if match:
            return match.group(1).replace("-", "/")

    match = re.search(r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b", texto)
    if match:
        return match.group(1).replace("-", "/")

    return ""


def limpiar_linea_paraclinico(linea):
    linea = normalizar_texto_para_reporte(linea)
    linea = re.sub(r"\s+", " ", linea).strip(" -\t")
    return linea


def normalizar_texto_para_reporte(texto):
    texto = str(texto or "")
    texto = unicodedata.normalize("NFKC", texto)
    texto = texto.replace("\r", "\n")
    texto = texto.replace("\xa0", " ")
    texto = texto.replace("ﬁ", "fi").replace("ﬂ", "fl")
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def compactar_espaciado_letras(texto):
    lineas_arregladas = []
    for linea in str(texto).splitlines():
        tokens = linea.split()
        if tokens:
            letras_sueltas = sum(1 for token in tokens if len(token) == 1 and token.isalpha())
            if letras_sueltas >= 5 and letras_sueltas / max(len(tokens), 1) > 0.5:
                reconstruida = ""
                buffer_letras = []
                for token in tokens:
                    if len(token) == 1 and token.isalpha():
                        buffer_letras.append(token)
                    else:
                        if buffer_letras:
                            reconstruida += "".join(buffer_letras) + " "
                            buffer_letras = []
                        reconstruida += token + " "
                if buffer_letras:
                    reconstruida += "".join(buffer_letras)
                linea = reconstruida.strip()
        lineas_arregladas.append(linea)
    return "\n".join(lineas_arregladas)


def organizar_texto_por_fechas(texto):
    texto = compactar_espaciado_letras(normalizar_texto_para_reporte(texto))
    if not texto:
        return ""

    texto = texto.upper()
    lineas = [limpiar_linea_paraclinico(linea) for linea in texto.splitlines()]
    lineas = [linea for linea in lineas if linea]

    patron_fecha = re.compile(r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b")
    bloques = []
    fecha_actual = None
    contenido_actual = []

    for linea in lineas:
        coincidencias = list(patron_fecha.finditer(linea))
        if coincidencias:
            for idx, match in enumerate(coincidencias):
                fecha = match.group(1).replace("-", "/")
                resto_inicio = match.end()
                resto_fin = coincidencias[idx + 1].start() if idx + 1 < len(coincidencias) else len(linea)
                resto = limpiar_linea_paraclinico(linea[resto_inicio:resto_fin])

                if fecha_actual is not None:
                    bloques.append((fecha_actual, contenido_actual))

                fecha_actual = fecha
                contenido_actual = [resto] if resto else []
        else:
            if fecha_actual is None:
                fecha_actual = "SIN FECHA"
                contenido_actual = []
            contenido_actual.append(linea)

    if fecha_actual is not None:
        bloques.append((fecha_actual, contenido_actual))

    bloques_agrupados = []
    mapa_fechas = {}
    for fecha, contenidos in bloques:
        if fecha not in mapa_fechas:
            mapa_fechas[fecha] = []
            bloques_agrupados.append((fecha, mapa_fechas[fecha]))
        for item in contenidos:
            item_limpio = limpiar_linea_paraclinico(item)
            if item_limpio:
                mapa_fechas[fecha].append(item_limpio)

    salida = []
    for fecha, contenidos in bloques_agrupados:
        if fecha != "SIN FECHA":
            salida.append(fecha)
        for item in contenidos:
            salida.append(item)
        salida.append("")

    return "\n".join(salida).strip()


def formatear_resumen_imagen(texto):
    texto = compactar_espaciado_letras(normalizar_texto_para_reporte(texto))
    if not texto:
        return ""

    fecha = extraer_fecha_principal(texto)
    lineas = [limpiar_linea_paraclinico(l).upper() for l in texto.splitlines() if limpiar_linea_paraclinico(l)]

    inicio = None
    for i, linea in enumerate(lineas):
        if any(
            marcador in linea
            for marcador in [
                "RADIOGRAFIA", "RADIOGRAFÍA", "TAC", "TOMOGRAFIA", "TOMOGRAFÍA",
                "ECOGRAFIA", "ECOGRAFÍA", "ULTRASON", "RESONANCIA", "RMN", "RM "
            ]
        ):
            inicio = i
            break

    if inicio is None:
        contenido = organizar_texto_por_fechas(texto)
        return contenido.upper()

    fin = len(lineas)
    for j in range(inicio + 1, len(lineas)):
        if any(
            stop in lineas[j]
            for stop in ["ATENTAMENTE", "GENERADO:", "PAGINA", "PÁGINA", "BACTERIOLOGA", "INTERPRETACION"]
        ):
            fin = j
            break

    bloque = lineas[inicio:fin]
    bloque = [b for b in bloque if b not in {"RESULTADO", "DATOS CLINICOS", "DATOS CLÍNICOS"}]
    contenido = " ".join(bloque)
    contenido = re.sub(r"\s+", " ", contenido).strip()
    contenido = re.sub(r"\s+[A-Z]$", "", contenido).strip()

    if fecha:
        return f"{fecha}\n{contenido}".strip()
    return contenido


def extraer_valor_numerico(linea, etiqueta):
    patron = re.escape(etiqueta) + r"\s+([<>]?\d+(?:[.,]\d+)?)"
    match = re.search(patron, linea)
    return match.group(1) if match else None


def linea_es_ruido_laboratorio(linea):
    ruido = [
        "PACIENTE", "DOCUMENTO ID", "FECHA DE NACIMIENTO", "DIRECCION", "DIRECCIÓN",
        "TELEFONO", "TELÉFONO", "SEXO", "FECHA DE INGRESO", "FECHA DE IMPRESION",
        "FECHA DE IMPRESIÓN", "SERVICIO", "EMPRESA", "MEDICO", "MÉDICO", "SEDE",
        "NO. INTERNO", "PETICION NO", "PETICIÓN NO", "PAGINA", "PÁGINA",
        "VALORES DE REFERENCIA", "UNIDADES", "FIRMA RESPONSABLE", "SEDE DE PROCESAMIENTO",
        "RAZON SOCIAL", "RAZÓN SOCIAL", "METODO:", "MÉTODO:", "BIBLIOGRAFIA:", "BIBLIOGRAFÍA:",
        "TITULO SIGNIFICATIVO", "TÍTULO SIGNIFICATIVO", "LOS TITULOS DE", "LOS TÍTULOS DE",
        "SE SUGIERE SEGUIMIENTO", "A PARTIR DEL", "MARCA REACTIVO", "SUSTRATO:",
        "LABORATORIO CLINICO", "LABORATORIO CLÍNICO", "PRELIMINAR", "A. M.", "P. M.",
        "CLINICA DE MARLY", "CLÍNICA DE MARLY", "COLSANITAS", "CAVELIER", "JOHANNA CORREDOR",
        "BOGOTA", "BOGOTÁ", "CHIA", "CHÍA", "COLOMBIA", "CUNDINAMARCA",
        "CARRERA ", "CALLE ", "AV.", "AV ", "VEREDA", "TEL.", "DIRECTOS:",
        "BACTERIOLOGA", "BACTERIÓLOGA", "HOMBRES:", "MUJERES:", "INMUNOENSAYO",
        "QUIMIOLUMINISCENCIA", "VALOR DE REFERENCIA", "VALORES DE REFERENCIA"
    ]
    return any(r in linea for r in ruido)


def linea_parece_nombre_examen(linea):
    linea = linea.strip()
    if not linea or len(linea) < 4 or len(linea) > 120:
        return False
    if linea_es_ruido_laboratorio(linea):
        return False
    if any(
        token in linea
        for token in [
            "RESULTADO", "VALOR DE REFERENCIA", "VALORES DE REFERENCIA", "METODO",
            "MÉTODO", "VALIDADO POR", "TECNICA", "TÉCNICA", "INTERVALO BIOLOGICO",
            "INTERVALO BIOLÓGICO", "UNIDADES", "PAGINA", "PÁGINA"
        ]
    ):
        return False
    if re.fullmatch(r"(PG/ML|NG/ML|MUI/ML|UUI/ML|U/ML|UI/ML|UI/L|%|UI)", linea):
        return False
    if re.fullmatch(r"[0-9./: %-]+", linea):
        return False
    if linea in {"COAGULACION", "COAGULACIÓN", "INMUNOLOGIA IV", "INMUNOLOGÍA IV", "INMUNOLOGIA E INFECCIOSAS ( SUERO)", "QUIMICA SANGUINEA (SUERO-QUIMICA SECA)"}:
        return False
    letras = sum(ch.isalpha() for ch in linea)
    return letras >= 4


def formatear_estudio_generico(linea):
    linea = re.sub(r"\s+", " ", linea).strip(" .")
    linea = linea.replace("%ACTIVIDAD", "% ACTIVIDAD")
    patron = re.match(
        r"^([A-ZÁÉÍÓÚÜÑ0-9 #/%()\-]+?)\s+((?:NEGATIVO|POSITIVO|NO REACTIVO|REACTIVO|[<>]?\d+(?:[.,]\d+)?(?:\s*%[A-ZÁÉÍÓÚÜÑ]*)?))\s*(.*)$",
        linea
    )
    if not patron:
        return ""

    nombre = patron.group(1).strip(" :")
    resultado = patron.group(2).strip()
    cola = patron.group(3).strip()
    cola = re.sub(r"^(%ACTIVIDAD|%|ACTIVIDAD)\s*", "", cola).strip()

    if cola:
        match_vn = re.match(r"([<>]?\d+(?:[.,]\d+)?)\s+([<>]?\d+(?:[.,]\d+)?)", cola)
        if match_vn:
            vn = f"VN {match_vn.group(1)} {match_vn.group(2)}"
            return f"{nombre} {resultado} {vn}".strip()
    return f"{nombre} {resultado}".strip()


def extraer_estudios_ocr(lineas):
    estudios = []
    exam_actual = None
    i = 0

    while i < len(lineas):
        linea = lineas[i]
        if not linea:
            i += 1
            continue

        if linea_parece_nombre_examen(linea):
            candidato = linea
            if i + 1 < len(lineas):
                siguiente = lineas[i + 1]
                if re.fullmatch(r"(PG/ML|NG/ML|MUI/ML|UUI/ML|U/ML|UI/ML|UI/L|%|UI)", siguiente):
                    i += 1
                    siguiente = lineas[i + 1] if i + 1 < len(lineas) else ""
                if (
                    linea_parece_nombre_examen(siguiente)
                    and normalizar_texto(siguiente) != normalizar_texto(candidato)
                    and len(siguiente) > len(candidato)
                ):
                    candidato = siguiente
                    i += 1
            exam_actual = candidato
            i += 1
            continue

        if exam_actual:
            linea_upper = linea.upper()
            valor = ""
            if "RESULTADO:" in linea_upper:
                valor = linea_upper.split("RESULTADO:", 1)[1].strip()
            elif re.fullmatch(r"(NEGATIVO|POSITIVO|VER ANEXO|NO REACTIVO|REACTIVO)(\s+[0-9.,]+)?", linea_upper):
                valor = linea_upper
            elif re.fullmatch(r"[<>]?\d+(?:[.,]\d+)?", linea_upper):
                valor = linea_upper

            if valor:
                estudio = f"{exam_actual} {valor}".strip()
                estudio_upper = estudio.upper()
                if any(
                    ruido in estudio_upper
                    for ruido in [
                        "BACTERIOLOG", "ANTICONCEPTIVOS", "SEMANAS DE GESTACION",
                        "SEMANAS DE GESTACIÓN", "MANTILLA", "CASTANEDA", "PRIETO",
                        "ALVAREZ", "30349355"
                    ]
                ):
                    exam_actual = None
                    i += 1
                    continue
                if estudio not in estudios:
                    estudios.append(estudio)
                exam_actual = None
                i += 1
                continue

        i += 1

    return estudios


def formatear_resumen_paraclinico(texto):
    texto = compactar_espaciado_letras(normalizar_texto_para_reporte(texto))
    if not texto:
        return ""

    fecha = extraer_fecha_principal(texto)
    upper = texto.upper()
    lineas = [limpiar_linea_paraclinico(l).upper() for l in upper.splitlines() if limpiar_linea_paraclinico(l)]

    salida = []
    if fecha:
        salida.append(fecha)

    pruebas_rapidas = []
    for linea in lineas:
        if "ANTIGENO" in linea or "ANTÍGENO" in linea:
            if "NEGAT" in linea:
                pruebas_rapidas.append(re.sub(r"\s+", " ", linea).replace("NEGATIVO", "NEGATIVO"))
            elif "POSIT" in linea:
                pruebas_rapidas.append(re.sub(r"\s+", " ", linea))
    salida.extend(dict.fromkeys(pruebas_rapidas))

    hemograma_map = [
        ("RECUENTO DE LEUCOCITOS", "LEUCOS"),
        ("% NEUTROFILOS", "NEU"),
        ("% NEUTRÓFILOS", "NEU"),
        ("% LINFOCITOS", "LINF"),
        ("HEMOGLOBINA", "HB"),
        ("HEMATOCRITO", "HTO"),
        ("RECUENTO DE PLAQUETAS", "PLAQ"),
        ("VOLUMEN CORPUSCULAR MEDIO", "VCM"),
        ("# NEUTROFILOS", "NEUTRO"),
        ("# NEUTRÓFILOS", "NEUTRO"),
        ("# LINFOCITOS", "L"),
    ]
    hemograma_items = []
    for origen, destino in hemograma_map:
        for linea in lineas:
            if linea.startswith(origen):
                valor = extraer_valor_numerico(linea, origen)
                if valor:
                    sufijo = "%" if origen.startswith("% ") else ""
                    hemograma_items.append(f"{destino} {valor}{sufijo}")
                break
    if hemograma_items:
        salida.append("HEMOGRAMA: " + ", ".join(hemograma_items))

    quimica_map = [
        ("VSG", "VSG"),
        ("PROTEINA C REACTIVA", "PCR"),
        ("PCR", "PCR"),
        ("ALBUMINA", "ALBÚMINA"),
        ("ALBÚMINA", "ALBÚMINA"),
        ("ALT", "ALT"),
        ("AST", "AST"),
        ("TGP", "TGP"),
        ("TGO", "TGO"),
        ("CREATININA", "CREATININA"),
        ("BUN", "BUN"),
        ("AMILASA", "AMILASA"),
        ("BILIRRUBINA TOTAL", "BT"),
        ("BILIRRUBINA DIRECTA", "BD"),
        ("BILIRRUBINA INDIRECTA", "BI"),
        ("CK", "CK"),
    ]

    grupos_quimica = []
    vistos = set()
    for origen, destino in quimica_map:
        for linea in lineas:
            if linea.startswith(origen) and (origen, destino) not in vistos:
                valor = extraer_valor_numerico(linea, origen)
                if valor:
                    grupos_quimica.append(f"{destino} {valor}")
                    vistos.add((origen, destino))
                break
    if grupos_quimica:
        salida.append(", ".join(grupos_quimica))

    uro_lineas = []
    capturando_uro = False
    for linea in lineas:
        if "UROANALISIS" in linea or "UROANÁLISIS" in linea:
            capturando_uro = True
        if capturando_uro:
            if any(stop in linea for stop in ["BACTERIOLOGA", "BACTERIÓLOGA", "PAGINA", "PÁGINA", "FECHA VALIDACION"]):
                break
            uro_lineas.append(linea)
    if uro_lineas:
        uro_texto = re.sub(r"\s+", " ", " ".join(uro_lineas)).strip()
        salida.append(uro_texto)

    if len(salida) <= 1:
        genericos = []
        en_area_examenes = False
        for linea in lineas:
            if ("EXAMEN" in linea or "EXÁMEN" in linea) and "RESULTADO" in linea:
                en_area_examenes = True
                continue
            if not en_area_examenes:
                continue
            if linea_es_ruido_laboratorio(linea):
                continue
            if linea in {"COAGULACION", "COAGULACIÓN", "INMUNOLOGIA IV", "INMUNOLOGÍA IV"}:
                continue
            if "FECHA VALIDACION" in linea or "FECHA VALIDACIÓN" in linea:
                continue
            if re.fullmatch(r"\d{1,2}/\d{1,2}/\d{2,4}", linea):
                continue
            estudio = formatear_estudio_generico(linea)
            if estudio:
                genericos.append(estudio)

        genericos = list(dict.fromkeys(genericos))
        if genericos:
            salida.extend(genericos)

    if len(salida) <= 1:
        estudios_ocr = extraer_estudios_ocr(lineas)
        if estudios_ocr:
            salida.extend(estudios_ocr)

    if len(salida) > 1:
        fecha_principal = salida[0]
        resto = [s for s in salida[1:] if s and not re.fullmatch(r"\d{1,2}/\d{1,2}/\d{2,4}", s)]
        return "\n".join([fecha_principal] + resto).strip()

    if len(salida) <= 1:
        return organizar_texto_por_fechas(texto).upper()


def organizar_pdf_segun_tipo(texto, tipo):
    if texto == "__PDF_ESCANEADO_SIN_TEXTO__":
        if tipo == "imagenes":
            return "PDF ESCANEADO SIN TEXTO EXTRAÍBLE. PARA ORGANIZAR AUTOMÁTICAMENTE ESTE REPORTE IMAGENOLÓGICO SE NECESITA OCR."
        return "PDF ESCANEADO SIN TEXTO EXTRAÍBLE. PARA ORGANIZAR AUTOMÁTICAMENTE ESTOS LABORATORIOS SE NECESITA OCR."
    if tipo == "imagenes":
        return formatear_resumen_imagen(texto)
    return formatear_resumen_paraclinico(texto)


def actualizar_texto_extraido(key_texto, key_auto, key_sig, pdf_files, tipo):
    if not pdf_files:
        return

    firma = "||".join(f"{pdf_file.name}|{pdf_file.size}" for pdf_file in pdf_files)
    if st.session_state.get(key_sig) == firma:
        return

    bloques = []
    for pdf_file in pdf_files:
        texto_pdf = extraer_texto_pdf(pdf_file)
        texto_organizado = organizar_pdf_segun_tipo(texto_pdf, tipo)
        if texto_organizado:
            bloques.append(texto_organizado)

    texto_organizado = "\n\n".join(bloques).strip()

    if not texto_organizado:
        return

    texto_actual = st.session_state.get(key_texto, "")
    texto_auto_previo = st.session_state.get(key_auto, "")
    if not texto_actual or texto_actual == texto_auto_previo:
        st.session_state[key_texto] = texto_organizado

    st.session_state[key_auto] = texto_organizado
    st.session_state[key_sig] = firma


def texto_a_html(texto):
    if not texto:
        return ""
    return "<br>".join(escape(str(texto)).splitlines())


def contenido_seccion_html(texto, color):
    if not texto:
        return f'<p style="margin:0 0 10px 0; color:{color};">&nbsp;</p>'

    lineas = str(texto).splitlines()
    partes = []
    for linea in lineas:
        linea_html = escape(linea) if linea.strip() else "&nbsp;"
        partes.append(
            f'<p style="margin:0 0 8px 0; line-height:1.5; color:{color};">{linea_html}</p>'
        )
    return "".join(partes)


def render_informe_html(titulo, secciones, texto_copiar):
    bloques_display = []
    bloques_copy = []
    for encabezado, contenido in secciones:
        bloques_display.append(
            f"""
            <section style="margin-top:18px;">
                <div style="font-weight:700; margin-bottom:8px; color:#ffffff;">{escape(encabezado)}</div>
                <div>{contenido_seccion_html(contenido, "#f5f5f5")}</div>
            </section>
            """
        )
        bloques_copy.append(
            f"""
            <section style="margin-top:16px;">
                <div style="font-weight:700; margin:0 0 8px 0; color:#111111;">{escape(encabezado)}</div>
                <div>{contenido_seccion_html(contenido, "#111111")}</div>
            </section>
            """
        )

    contenido_html = f"""
    <div style="font-family:Arial, sans-serif; color:#111111;">
        <h1 style="font-size:20px; font-weight:800; text-align:center; margin:0 0 18px 0;">{escape(titulo)}</h1>
        {''.join(bloques_copy)}
    </div>
    """

    html = f"""
    <div style="border:1px solid rgba(250,250,250,.12); border-radius:14px; padding:18px 20px; background:#171923; color:#f5f5f5; font-family:Arial, sans-serif;">
        <div style="display:flex; justify-content:flex-end; margin-bottom:8px;">
            <button
                onclick='(async () => {{
                    const htmlContent = {json.dumps(contenido_html)};
                    const textContent = {json.dumps(texto_copiar)};
                    try {{
                        const item = new ClipboardItem({{
                            "text/html": new Blob([htmlContent], {{ type: "text/html" }}),
                            "text/plain": new Blob([textContent], {{ type: "text/plain" }})
                        }});
                        await navigator.clipboard.write([item]);
                        this.textContent = "Informe copiado";
                        setTimeout(() => this.textContent = "Copiar informe", 1500);
                    }} catch (err) {{
                        await navigator.clipboard.writeText(textContent);
                        this.textContent = "Copiado sin formato";
                        setTimeout(() => this.textContent = "Copiar informe", 1500);
                    }}
                }})()'
                style="background:#2b6cb0; color:white; border:none; border-radius:8px; padding:8px 12px; cursor:pointer; font-size:14px;"
            >
                Copiar informe
            </button>
        </div>
        <div style="font-size:24px; font-weight:800; text-align:center; margin-bottom:18px;">{escape(titulo)}</div>
        {''.join(bloques_display)}
    </div>
    """
    components.html(html, height=1200, scrolling=True)


def formatear_numero_clinico(valor, decimales=1):
    try:
        numero = round(float(valor), decimales)
    except Exception:
        return ""

    texto = f"{numero:.{decimales}f}"
    if "." in texto:
        texto = texto.rstrip("0").rstrip(".")
    return texto


def extraer_codigo_cie10(diagnostico):
    if not diagnostico:
        return ""
    match = re.match(r"\s*([A-Z]\d{2}(?:\.\d+)?)", str(diagnostico).strip(), flags=re.IGNORECASE)
    return match.group(1).upper() if match else ""


def calcular_dosis_mg(peso, mg_kg, max_mg=None):
    if not peso or peso <= 0:
        return None
    dosis = peso * mg_kg
    if max_mg is not None:
        dosis = min(dosis, max_mg)
    return round(dosis, 2)


def calcular_volumen_ml(dosis_mg, concentracion_mg_en_5ml):
    if not dosis_mg or not concentracion_mg_en_5ml:
        return None
    mg_por_ml = concentracion_mg_en_5ml / 5
    if mg_por_ml <= 0:
        return None
    return round(dosis_mg / mg_por_ml, 2)


def calcular_liquido_mantenimiento_holliday(peso):
    if not peso or peso <= 0:
        return None
    if peso <= 10:
        ml_hora = (peso * 100) / 24
    elif peso <= 20:
        ml_hora = ((1000 + (peso - 10) * 50) / 24)
    else:
        ml_hora = ((1500 + (peso - 20) * 20) / 24)
    return round(ml_hora, 1)


def calcular_liquido_superficie_corporal(peso, talla):
    if not peso or not talla or peso <= 0 or talla <= 0:
        return None, None
    sc = calcular_sc(peso, talla)
    ml_dia = round(sc * 1500, 1)
    ml_hora = round(ml_dia / 24, 1)
    return round(sc, 2), ml_hora


def calcular_parkland(peso, scq_pct):
    if not peso or peso <= 0 or not scq_pct or scq_pct <= 0:
        return None, None, None
    total_24h = round(4 * peso * scq_pct, 1)
    primeras_8h = round(total_24h / 2, 1)
    siguientes_16h = round(total_24h / 2, 1)
    return total_24h, primeras_8h, siguientes_16h


def calcular_galveston(peso, talla, scq_pct):
    if not peso or peso <= 0 or not talla or talla <= 0 or not scq_pct or scq_pct <= 0:
        return None, None
    sc = calcular_sc(peso, talla)
    scq_m2 = sc * (scq_pct / 100)
    total_24h = round((5000 * scq_m2) + (2000 * sc), 1)
    ml_hora = round(total_24h / 24, 1)
    return total_24h, ml_hora


def construir_contexto_plan_medicacion(peso):
    dosis_cfg = cargar_dosis_medicacion()

    acet_cfg = dosis_cfg["ACETAMINOFEN"]
    acet_mg = calcular_dosis_mg(peso, acet_cfg["mg_kg_dosis"], max_mg=acet_cfg.get("max_mg")) if peso else None
    if acet_mg:
        linea_acetaminofen = (
            f"- ACETAMINOFEN: {formatear_numero_clinico(acet_mg, 0)} MG {acet_cfg.get('via', 'VO')} "
            f"CADA {int(acet_cfg.get('intervalo_horas', 6))} HORAS ({acet_cfg.get('indicacion', '')})"
        )
    else:
        linea_acetaminofen = "- ACETAMINOFEN VO SEGUN PESO (SI FIEBRE O DOLOR)"

    dip_cfg = dosis_cfg["DIPIRONA"]
    dip_mg = calcular_dosis_mg(peso, dip_cfg["mg_kg_dosis"], max_mg=dip_cfg.get("max_mg")) if peso else None
    if dip_mg:
        linea_dipirona = (
            f"- DIPIRONA: {formatear_numero_clinico(dip_mg, 0)} MG {dip_cfg.get('via', 'IV')} "
            f"CADA {int(dip_cfg.get('intervalo_horas', 8))} HORAS ({dip_cfg.get('indicacion', '')})"
        )
    else:
        linea_dipirona = "- DIPIRONA IV SEGUN PESO (SI FIEBRE O DOLOR NO CEDE CON ACETAMINOFEN)"

    dexa_mg = calcular_dosis_mg(peso, 0.6, max_mg=10) if peso else None
    if dexa_mg:
        linea_dexametasona_laringitis = f"- DEXAMETASONA {formatear_numero_clinico(dexa_mg, 1)} MG IV DOSIS UNICA"
    else:
        linea_dexametasona_laringitis = "- DEXAMETASONA 0.6 MG/KG IV DOSIS UNICA"

    return {
        "LINEA_ACETAMINOFEN": linea_acetaminofen,
        "LINEA_DIPIRONA": linea_dipirona,
        "LINEA_DEXAMETASONA_LARINGITIS": linea_dexametasona_laringitis,
    }


def renderizar_plan_editable(texto_plan, peso):
    contexto = construir_contexto_plan_medicacion(peso)
    try:
        return texto_plan.format_map(SafeFormatDict(contexto))
    except Exception:
        return texto_plan


def generar_plan_base_ordenado(peso, edad_meses, incluir_analgesia=True):
    contexto_medicacion = construir_contexto_plan_medicacion(peso)
    lineas = [
        "- HOSPITALIZACION PEDIATRICA",
        "- DIETA NORMAL ACORDE A LA EDAD",
        "- CATETER SELLADO",
    ]

    liquidos_ml_hora = calcular_liquido_mantenimiento_holliday(peso)
    if liquidos_ml_hora:
        lineas.append(f"- LACTATO DE RINGER A {formatear_numero_clinico(liquidos_ml_hora, 1)} CC/HORA IV")
    else:
        lineas.append("- LACTATO DE RINGER IV SEGUN REQUERIMIENTO CLINICO")

    if incluir_analgesia and peso and peso > 0:
        dipi_cfg = cargar_dosis_medicacion()["DIPIRONA"]
        ibup_mg = calcular_dosis_mg(peso, 10, max_mg=400)

        lineas.append(contexto_medicacion["LINEA_ACETAMINOFEN"])
        lineas.append(contexto_medicacion["LINEA_DIPIRONA"])

        if edad_meses is not None and edad_meses >= 6 and ibup_mg:
            lineas.append(
                f"- IBUPROFENO: {formatear_numero_clinico(ibup_mg, 0)} MG VO CADA 8 HORAS "
                f"(SI DOLOR O FIEBRE)"
            )
    else:
        lineas.append(contexto_medicacion["LINEA_ACETAMINOFEN"])
        lineas.append(contexto_medicacion["LINEA_DIPIRONA"])

    lineas.extend([
        "- SS PARACLINICOS DE EXTENSIÓN",
        "- CONTROL DE LÍQUIDOS ADMINISTRADOS - ELIMINADOS",
        "- CONTROL DE SIGNOS VITALES, AVISAR CAMBIOS.",
    ])
    return lineas


def generar_resumen_dosis(diagnostico, peso, talla, edad_meses, scq_pct=0):
    lineas = []
    if not peso or peso <= 0:
        return "INGRESE PESO PARA CALCULAR DOSIS AUTOMÁTICAS."

    dosis_cfg = cargar_dosis_medicacion()
    acet_cfg = dosis_cfg["ACETAMINOFEN"]
    dip_cfg = dosis_cfg["DIPIRONA"]
    acet_mg = calcular_dosis_mg(peso, acet_cfg["mg_kg_dosis"], max_mg=acet_cfg.get("max_mg"))
    dipi_mg = calcular_dosis_mg(peso, dip_cfg["mg_kg_dosis"], max_mg=dip_cfg.get("max_mg"))
    ibup_mg = calcular_dosis_mg(peso, 10, max_mg=400)
    holliday_ml_h = calcular_liquido_mantenimiento_holliday(peso)
    sc, sc_ml_h = calcular_liquido_superficie_corporal(peso, talla)
    parkland_total, parkland_8h, parkland_16h = calcular_parkland(peso, scq_pct)
    galveston_total, galveston_ml_h = calcular_galveston(peso, talla, scq_pct)

    if holliday_ml_h:
        lineas.append(f"- LÍQUIDOS DE MANTENIMIENTO HOLLIDAY: {formatear_numero_clinico(holliday_ml_h, 1)} CC/HORA")
    if sc is not None and sc_ml_h is not None:
        lineas.append(
            f"- SUPERFICIE CORPORAL: {formatear_numero_clinico(sc, 2)} M2 | MANTENIMIENTO POR SC (1500 ML/M2/DÍA): {formatear_numero_clinico(sc_ml_h, 1)} CC/HORA"
        )
    if parkland_total is not None:
        lineas.append(
            f"- PARKLAND ({formatear_numero_clinico(scq_pct,1)}% SCQ): {formatear_numero_clinico(parkland_total, 1)} ML/24H | "
            f"{formatear_numero_clinico(parkland_8h, 1)} ML EN 8H Y {formatear_numero_clinico(parkland_16h, 1)} ML EN 16H"
        )
    if galveston_total is not None and galveston_ml_h is not None:
        lineas.append(
            f"- GALVESTON ({formatear_numero_clinico(scq_pct,1)}% SCQ): {formatear_numero_clinico(galveston_total, 1)} ML/24H = "
            f"{formatear_numero_clinico(galveston_ml_h, 1)} CC/HORA"
        )
    if acet_mg:
        lineas.append(
            f"- ACETAMINOFEN {formatear_numero_clinico(acet_cfg['mg_kg_dosis'], 1)} MG/KG/DOSIS: "
            f"{formatear_numero_clinico(acet_mg, 0)} MG CADA {int(acet_cfg.get('intervalo_horas', 6))} HORAS"
        )
    if dipi_mg:
        lineas.append(
            f"- DIPIRONA {formatear_numero_clinico(dip_cfg['mg_kg_dosis'], 1)} MG/KG/DOSIS: "
            f"{formatear_numero_clinico(dipi_mg, 0)} MG {dip_cfg.get('via', 'IV')} CADA {int(dip_cfg.get('intervalo_horas', 8))} HORAS"
        )
    if edad_meses is not None and edad_meses >= 6 and ibup_mg:
        lineas.append(f"- IBUPROFENO 10 MG/KG/DOSIS: {formatear_numero_clinico(ibup_mg, 0)} MG")

    codigo = extraer_codigo_cie10(diagnostico)
    if codigo.startswith("J18"):
        ceftriaxona_dia = calcular_dosis_mg(peso, 50, max_mg=2000)
        if ceftriaxona_dia:
            lineas.append(f"- CEFTRIAXONA 50 MG/KG/DÍA: {formatear_numero_clinico(ceftriaxona_dia, 0)} MG IV CADA 24 HORAS")
    elif codigo.startswith("H66"):
        amoxi_dia = calcular_dosis_mg(peso, 90, max_mg=4000)
        amoxi_dosis = round(amoxi_dia / 2, 2) if amoxi_dia else None
        if amoxi_dosis:
            lineas.append(f"- AMOXICILINA 90 MG/KG/DÍA: {formatear_numero_clinico(amoxi_dosis, 0)} MG CADA 12 HORAS")
    elif codigo.startswith("J03"):
        amoxi_dia = calcular_dosis_mg(peso, 50, max_mg=3000)
        amoxi_dosis = round(amoxi_dia / 2, 2) if amoxi_dia else None
        if amoxi_dosis:
            lineas.append(f"- AMOXICILINA 50 MG/KG/DÍA: {formatear_numero_clinico(amoxi_dosis, 0)} MG CADA 12 HORAS")
    elif codigo.startswith("J05") or codigo.startswith("J04"):
        dexa_mg = calcular_dosis_mg(peso, 0.6, max_mg=10)
        adrenalina_mg = min(round(peso * 0.5, 2), 5) if peso else None
        if dexa_mg:
            lineas.append(f"- DEXAMETASONA 0.6 MG/KG DOSIS ÚNICA: {formatear_numero_clinico(dexa_mg, 1)} MG")
        if adrenalina_mg is not None:
            lineas.append(f"- ADRENALINA NEBULIZADA 0.5 MG/KG (MÁX 5 MG): {formatear_numero_clinico(adrenalina_mg, 1)} MG")
    elif codigo.startswith("A09"):
        ondasetron_mg = calcular_dosis_mg(peso, 0.15, max_mg=8)
        if ondasetron_mg:
            lineas.append(f"- ONDANSETRON 0.15 MG/KG/DOSIS: {formatear_numero_clinico(ondasetron_mg, 1)} MG")
    elif codigo.startswith("J21"):
        salbutamol_mg = calcular_dosis_mg(peso, 0.15, max_mg=5)
        if salbutamol_mg:
            lineas.append(f"- SALBUTAMOL 0.15 MG/KG/NEBULIZACIÓN: {formatear_numero_clinico(salbutamol_mg, 1)} MG SI ESTÁ INDICADO.")

    if not lineas:
        return "SIN CÁLCULOS DISPONIBLES PARA EL DIAGNÓSTICO ACTUAL."
    return "\n".join(lineas)


def generar_plan_sugerido(diagnostico, peso, edad_meses):
    codigo = extraer_codigo_cie10(diagnostico)
    if not codigo and not es_diagnostico_laringitis(diagnostico):
        return "\n".join(generar_plan_base_ordenado(peso, edad_meses))

    planes_patologia = cargar_planes_patologia()
    clave_plan = obtener_clave_plan_patologia(codigo)
    plan_patologia_editable = planes_patologia.get(clave_plan, "").strip() if clave_plan else ""

    if es_diagnostico_laringitis(diagnostico):
        if plan_patologia_editable:
            return renderizar_plan_editable(plan_patologia_editable, peso)
        return renderizar_plan_editable(PLANES_PATOLOGIA_DEFAULTS["J05_J04"], peso)

    if plan_patologia_editable:
        return renderizar_plan_editable(plan_patologia_editable, peso)

    lineas = generar_plan_base_ordenado(peso, edad_meses)
    return "\n".join(lineas)


def generar_docx_informe(titulo, secciones):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Arial"
    estilo_normal.font.size = Pt(11)

    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_p.add_run(titulo)
    titulo_run.bold = True
    titulo_run.font.size = Pt(14)
    titulo_p.paragraph_format.space_after = Pt(8)
    titulo_p.paragraph_format.line_spacing = 1

    for index, (encabezado, contenido) in enumerate(secciones):
        p_head = doc.add_paragraph()
        run_head = p_head.add_run(encabezado)
        run_head.bold = True
        run_head.font.name = "Arial"
        run_head.font.size = Pt(11)
        p_head.paragraph_format.space_before = Pt(8 if index > 0 else 0)
        p_head.paragraph_format.space_after = Pt(0)
        p_head.paragraph_format.line_spacing = 1

        p_body = doc.add_paragraph()
        p_body.paragraph_format.space_before = Pt(0)
        p_body.paragraph_format.space_after = Pt(0)
        p_body.paragraph_format.line_spacing = 1

        lineas = str(contenido).splitlines() if contenido else [""]
        for line_index, linea in enumerate(lineas):
            run = p_body.add_run(linea)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            if line_index < len(lineas) - 1:
                run.add_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def cargar_historias_guardadas():
    if not HISTORIAS_PATH.exists():
        return []

    historias = []
    with HISTORIAS_PATH.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                historias.append(json.loads(line))
            except json.JSONDecodeError:
                continue

    return list(reversed(historias))


def render():
    if st.session_state.pop("_limpiar_formulario", False):
        limpiar_formulario()
        st.rerun()

    for key, value in FORM_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = value

    titulo_historia = st.session_state.get(
        "tipo_historia_clinica_ped_urg",
        "HISTORIA CLINICA DE INGRESO A URGENCIAS PEDIATRICAS"
    )

    st.header(f"📌 {titulo_historia}")

    col1, col2 = st.columns(2)

    with col1:
        nombre = st.text_input("Nombres y apellidos", key="nombre_1")
        tipo_documento = st.selectbox(
            "Tipo de documento",
            ["NV", "RC", "TI", "PEP", "PS", "OTRO"],
            index=None,
            placeholder="Seleccione tipo de documento",
            key="tipo_documento_1"
        )
        fecha_nacimiento = st.date_input(
            "Fecha de nacimiento",
            value=None,
            min_value=date(1900, 1, 1),
            max_value=date.today(),
            format="DD/MM/YYYY",
            key="fecha_1"
        )
        eps = st.text_input("EPS", key="eps_1")

    with col2:
        sexo = st.selectbox(
            "Sexo",
            ["Masculino", "Femenino"],
            index=None,
            placeholder="Seleccione sexo",
            key="sexo_1"
        )
        documento = st.text_input("Documento", key="documento_1")
        informante = st.text_input("Informante (parentesco - edad - ocupacion)", key="informante_1")
        proveniente = st.text_input("Proveniente", key="proveniente_1")

    z_pe = z_te = z_imc = z_pt = z_pc = None
    dx_nutricional = "NO EVALUADO"

    años = meses = 0
    grupo = ""
    sexo_data = {"paciente": "paciente"}

    if fecha_nacimiento:
        años, meses, dias = calcular_edad(fecha_nacimiento)
        grupo = grupo_etario(años)
        sexo_data = sexo_texto(sexo)

        st.info(f"Edad: {años} años, {meses} meses, {dias} días")

    motivo = st.text_area("Motivo de consulta", key="motivo_1")
    enfermedad_input = st.text_area("Detalles enfermedad actual", key="enfermedad_1")

    st.subheader("Antecedentes")

    antecedentes = st.text_area(
        "Antecedentes generales",
        key="antecedentes_1",
        height=250
    )

    neuro_editable = ""
    if fecha_nacimiento:
        neuro = obtener_neurodesarrollo(años, meses)

        st.subheader("Neurodesarrollo")

        if "neuro_prev" not in st.session_state or st.session_state["neuro_prev"] != neuro:
            st.session_state["neuro_1"] = neuro
            st.session_state["neuro_prev"] = neuro

        neuro_editable = st.text_area("Neurodesarrollo", key="neuro_1", height=200)

    st.subheader("Revisión por sistemas")
    revision = st.text_area(
        "Revisión",
        key="revision"
    )

    st.subheader("Signos vitales")

    col1, col2 = st.columns(2)

    with col1:
        ta = st.text_input("TA (mmHg)", key="ta")
        fc = st.number_input("FC (lpm)", min_value=0.0, key="fc")
        fr = st.number_input("FR (rpm)", min_value=0.0, key="fr")

    with col2:
        sat = st.number_input("SpO2 (%)", min_value=0.0, key="sat")
        temp = st.number_input("Temperatura (°C)", min_value=0.0, key="temp")

    peso = st.number_input("Peso (kg)", min_value=0.0, key="peso")
    talla = st.number_input("Talla (cm)", min_value=0.0, key="talla")
    pc = st.number_input("Perímetro cefálico (cm)", min_value=0.0, key="pc")
    scq_pct = st.number_input("Superficie corporal quemada (%)", min_value=0.0, max_value=100.0, key="scq_pct")

    if fecha_nacimiento and peso > 0 and talla > 0:

        edad_meses = edad_en_meses(fecha_nacimiento)
        sexo_oms = 1 if sexo == "Masculino" else 2

        imc = calcular_imc(peso, talla)

        z_pe = zscore_peso_edad(peso, edad_meses, sexo_oms)
        z_te = zscore_talla_edad(talla, edad_meses, sexo_oms)
        z_imc = zscore_imc_edad(imc, edad_meses, sexo_oms)
        z_pt = zscore_peso_talla(peso, talla, sexo_oms, edad_meses)
        z_pc = zscore_pc_edad(pc, edad_meses, sexo_oms)

        st.subheader("OMS Automático 🔥")
        st.write(f"P/E Z: {z_pe}")
        st.write(f"T/E Z: {z_te}")
        st.write(f"P/T Z: {z_pt}")
        st.write(f"IMC/E Z: {z_imc}")
        st.write(f"PC/E Z: {z_pc}")

        # 🔥 DIAGNÓSTICO NUTRICIONAL
        if edad_meses < 60:
            dx_nutricional = diagnostico_menor_5(z_pt, z_pe, z_te, z_pc)
        else:
            dx_nutricional = diagnostico_mayor_5(z_imc, z_te)

    # =========================
    # EXAMEN FÍSICO
    # =========================
    st.subheader("Examen físico")

    examen = st.text_area("Examen físico", key="examen", height=300)

    # =========================
    # PARACLÍNICOS / IMÁGENES
    # =========================
    st.subheader("Paraclínicos")
    col_pdf_1, col_pdf_2 = st.columns(2)

    with col_pdf_1:
        pdf_paraclinicos = st.file_uploader(
            "Subir PDF de laboratorios",
            type=["pdf"],
            key="pdf_paraclinicos_uploader",
            accept_multiple_files=True,
            help="Carga un PDF de laboratorio. La app extrae el texto, lo pone en MAYÚSCULA y lo organiza por fechas."
        )
        actualizar_texto_extraido(
            "paraclinicos_texto",
            "paraclinicos_auto",
            "paraclinicos_pdf_sig",
            pdf_paraclinicos,
            "paraclinicos"
        )

    with col_pdf_2:
        pdf_imagenes = st.file_uploader(
            "Subir PDF de imágenes",
            type=["pdf"],
            key="pdf_imagenes_uploader",
            accept_multiple_files=True,
            help="Carga un PDF de reporte imagenológico. La app extrae el texto, lo pone en MAYÚSCULA y lo organiza por fechas."
        )
        actualizar_texto_extraido(
            "imagenes_texto",
            "imagenes_auto",
            "imagenes_pdf_sig",
            pdf_imagenes,
            "imagenes"
        )

    paraclinicos_texto = st.text_area(
        "Laboratorios",
        key="paraclinicos_texto",
        height=220
    )

    imagenes_texto = st.text_area(
        "Imágenes",
        key="imagenes_texto",
        height=220
    )

    # =========================
    # ANÁLISIS
    # =========================
    if fecha_nacimiento:
        años, meses, dias = calcular_edad(fecha_nacimiento)
        edad_texto = f"{años} AÑOS" if años > 0 else f"{meses} MESES"
    else:
        edad_texto = ""

    texto_enfermedad = enfermedad_input.upper() if enfermedad_input else ""
    sexo_texto_analisis = sexo.upper() if sexo else ""

    enfermedad_auto = f"PACIENTE {sexo_texto_analisis} {grupo.upper()} DE {edad_texto}, QUIEN CONSULTA POR {texto_enfermedad}"

    analisis_default = f"""{enfermedad_auto}. AL INGRESO PACIENTE QUE LUCE EN ACEPTABLES CONDICIONES GENERALES, BUEN ESTADO DE HIDRATACIÓN, HEMODINÁMICAMENTE ESTABLE, BUEN PATRÓN RESPIRATORIO, SIN REQUERIMIENTO DE O2 SUPLEMENTARIO, SIN SIGNOS DE ALARMA ABDOMINAL, SIN DISTERMIAS, NO ASPECTO TÓXICO, SIN EDEMAS, SIN DÉFICIT NEUROLÓGICO, PIEL BIEN PERFUNDIDA SIN LESIONES. SE INDICA MANEJO... SE BRINDA INFORMACIÓN A FAMILIARES, SE ACLARAN DUDAS."""

    st.subheader("Análisis")

    # Carga el análisis por defecto una sola vez y solo lo refresca si el
    # contenido previo seguía siendo el automático.
    if "analisis" not in st.session_state:
        st.session_state["analisis"] = analisis_default
        st.session_state["analisis_base"] = analisis_default
    elif st.session_state.get("analisis_base") != analisis_default:
        if st.session_state.get("analisis") == st.session_state.get("analisis_base"):
            st.session_state["analisis"] = analisis_default
        st.session_state["analisis_base"] = analisis_default

    analisis = st.text_area("Análisis clínico", key="analisis", height=200)

    st.subheader("Diagnósticos")
    clasificacion_diagnostica = "CIE-10"
    cie10 = cargar_cie10()
    busqueda_cie10 = st.text_input(
        "Buscar CIE-10 por código o descripción",
        key="busqueda_cie10"
    ).strip()

    if busqueda_cie10:
        grupos_busqueda = construir_grupos_busqueda(busqueda_cie10)
        if grupos_busqueda:
            terminos = aplanar_grupos_busqueda(grupos_busqueda)
        else:
            terminos = list(expandir_terminos_busqueda(busqueda_cie10))

        cie10_filtrado = cie10.copy()
        if grupos_busqueda:
            cie10_filtrado_estricto = cie10_filtrado[
                cie10_filtrado.apply(lambda row: coincide_grupos(row, grupos_busqueda), axis=1)
            ]
            if not cie10_filtrado_estricto.empty:
                cie10_filtrado = cie10_filtrado_estricto
        if cie10_filtrado.empty:
            cie10_filtrado = cie10.iloc[0:0].copy()
        else:
            cie10_filtrado["score_busqueda"] = cie10_filtrado.apply(
                lambda row: puntuar_diagnostico(row, terminos, grupos_busqueda),
                axis=1
            )
            cie10_filtrado = cie10_filtrado[cie10_filtrado["score_busqueda"] > 0]
            cie10_filtrado = cie10_filtrado.sort_values(
                by=["score_busqueda", "code_normalized"],
                ascending=[False, True]
            ).head(20)
    else:
        cie10_filtrado = cie10.iloc[0:0]

    if cie10_filtrado.empty:
        if busqueda_cie10:
            st.caption("No se encontraron diagnósticos CIE-10 con esa búsqueda.")
        diagnostico_seleccionado = ""
    else:
        cie10_filtrado = cie10_filtrado.copy()
        cie10_filtrado["description_es"] = cie10_filtrado["description"].map(traducir_cie10_descripcion)
        cie10_filtrado["label_es"] = cie10_filtrado["code"].astype(str) + " - " + cie10_filtrado["description_es"]
        with st.expander(f"Resultados de diagnóstico ({len(cie10_filtrado)})", expanded=False):
            diagnostico_seleccionado = st.selectbox(
                "Diagnóstico CIE-10",
                cie10_filtrado["label_es"].tolist(),
                index=None,
                placeholder="Seleccione un diagnóstico",
                key="dx_cie10"
            )

    observacion_diagnostico = st.text_area(
        "Observación diagnóstica",
        key="obs_dx",
        height=100
    )

    edad_meses_actual = edad_en_meses(fecha_nacimiento) if fecha_nacimiento else None
    diagnostico_plan = diagnostico_seleccionado or ""
    plan_sugerido = generar_plan_sugerido(diagnostico_plan, peso, edad_meses_actual)

    if "plan" not in st.session_state:
        st.session_state["plan"] = plan_sugerido
        st.session_state["plan_base"] = plan_sugerido
    elif st.session_state.get("plan_base") != plan_sugerido:
        if st.session_state.get("plan") == st.session_state.get("plan_base"):
            st.session_state["plan"] = plan_sugerido
        st.session_state["plan_base"] = plan_sugerido

    plan = st.text_area(
        "Plan",
        key="plan",
        height=200
    )

    # =========================
    # GENERAR / LIMPIAR
    # =========================
    col_btn_1, col_btn_2 = st.columns(2)
    generar_historia = col_btn_1.button("Generar Historia Clínica", use_container_width=True)
    col_btn_2.button(
        "Limpiar y empezar otra historia",
        use_container_width=True,
        on_click=solicitar_limpieza_formulario
    )

    if generar_historia:

        fecha_str = fecha_nacimiento.strftime("%d/%m/%Y") if fecha_nacimiento else ""
        diagnostico_final = diagnostico_seleccionado or ""
        paraclinicos_final = paraclinicos_texto.strip() if str(paraclinicos_texto).strip() else "NO HAY REPORTES"
        imagenes_final = imagenes_texto.strip() if str(imagenes_texto).strip() else "NO HAY REPORTES"

        historia = f"""
{titulo_historia.upper()}

DATOS DE IDENTIFICACIÓN:

NOMBRE: {nombre}
TIPO DE DOCUMENTO: {tipo_documento}
DOCUMENTO: {documento}
FECHA DE NACIMIENTO: {fecha_str}
INFORMANTE: {informante}
EPS: {eps}
PROVENIENTE: {proveniente}

MOTIVO DE CONSULTA:
"{motivo}"

ENFERMEDAD ACTUAL:
{enfermedad_auto}

REVISIÓN POR SISTEMAS:
{revision}

ANTECEDENTES:
{antecedentes}

NEURODESARROLLO:
{neuro_editable}

SIGNOS VITALES:
TA {ta} mmHg FC: {fc} lpm FR: {fr} rpm SpO2: {sat}% T: {temp} °C

ANTROPOMETRÍA:
PESO: {peso} kg TALLA: {talla} cm PC: {pc} cm
P/E Z: {z_pe}
T/E Z: {z_te}
P/T Z: {z_pt}
IMC/E Z: {z_imc}
PC/E Z: {z_pc}

EXAMEN FÍSICO:
{examen}

LABORATORIOS:
{paraclinicos_final}

IMÁGENES:
{imagenes_final}

ANÁLISIS:
{analisis}

CLASIFICACIÓN DIAGNÓSTICA:
{clasificacion_diagnostica}

DIAGNÓSTICOS:
{diagnostico_final}

OBSERVACIÓN DIAGNÓSTICA:
{observacion_diagnostico}

DIAGNÓSTICO NUTRICIONAL:
{dx_nutricional}

PLAN:
{plan}
"""

        fecha_guardado = datetime.now(BOGOTA_TZ).strftime("%Y-%m-%d %H:%M:%S")
        identificador = f"{fecha_guardado} | {nombre or 'SIN NOMBRE'} | {documento or 'SIN DOCUMENTO'}"

        guardar_historia({
            "id": identificador,
            "fecha_guardado": fecha_guardado,
            "nombre": nombre,
            "tipo_documento": tipo_documento,
            "documento": documento,
            "historia": historia.upper()
        })

        st.success("Historia clínica generada")
        secciones_informe = [
            ("DATOS DE IDENTIFICACIÓN", f"NOMBRES Y APELLIDOS: {nombre}\nTIPO DE DOCUMENTO: {tipo_documento}\nDOCUMENTO: {documento}\nFECHA DE NACIMIENTO: {fecha_str}\nINFORMANTE: {informante}\nEPS: {eps}\nPROVENIENTE: {proveniente}"),
            ("MOTIVO DE CONSULTA", motivo),
            ("ENFERMEDAD ACTUAL", enfermedad_auto),
            ("REVISIÓN POR SISTEMAS", revision),
            ("ANTECEDENTES", antecedentes),
            ("NEURODESARROLLO", neuro_editable),
            ("SIGNOS VITALES", f"TA {ta} mmHg FC: {fc} lpm FR: {fr} rpm SpO2: {sat}% T: {temp} °C"),
            ("ANTROPOMETRÍA", f"PESO: {peso} kg TALLA: {talla} cm PC: {pc} cm\nP/E Z: {z_pe}\nT/E Z: {z_te}\nP/T Z: {z_pt}\nIMC/E Z: {z_imc}\nPC/E Z: {z_pc}"),
            ("EXAMEN FÍSICO", examen),
            ("LABORATORIOS", paraclinicos_final),
            ("IMÁGENES", imagenes_final),
            ("ANÁLISIS", analisis),
            ("CLASIFICACIÓN DIAGNÓSTICA", clasificacion_diagnostica),
            ("DIAGNÓSTICOS", diagnostico_final),
            ("OBSERVACIÓN DIAGNÓSTICA", observacion_diagnostico),
            ("DIAGNÓSTICO NUTRICIONAL", dx_nutricional),
            ("PLAN", plan),
        ]
        docx_bytes = generar_docx_informe(titulo_historia.upper(), secciones_informe)
        st.download_button(
            "Descargar informe en Word",
            data=docx_bytes,
            file_name=f"{(nombre or 'historia').strip().replace(' ', '_')}_historia_clinica.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        render_informe_html(titulo_historia.upper(), secciones_informe, historia.upper())

    st.divider()
    with st.expander("Historias guardadas", expanded=False):
        historias_guardadas = cargar_historias_guardadas()

        if historias_guardadas:
            opciones_historias = [h["id"] for h in historias_guardadas]
            historia_consulta_id = st.selectbox(
                "Consultar historia previa",
                opciones_historias,
                key="historia_consulta_id"
            )

            historia_seleccionada = next(
                (h for h in historias_guardadas if h["id"] == historia_consulta_id),
                None
            )

            if historia_seleccionada:
                st.text_area(
                    "Informe guardado",
                    historia_seleccionada["historia"],
                    height=500,
                    key="historia_guardada_texto"
                )
        else:
            st.info("Aún no hay historias guardadas.")

    with st.expander("Planes de manejo por patología", expanded=False):
        planes_patologia = cargar_planes_patologia()
        labels_planes_patologia = cargar_labels_planes_patologia()

        st.markdown("**Agregar patología nueva**")
        nueva_clave_patologia = st.text_input(
            "Nombre corto o código interno de la patología",
            key="nueva_clave_patologia"
        )
        nuevo_label_patologia = st.text_input(
            "Nombre visible de la patología",
            key="nuevo_label_patologia"
        )
        nuevo_plan_patologia = st.text_area(
            "Plan para la nueva patología",
            key="nuevo_plan_patologia",
            height=180
        )

        if st.button("Agregar patología a la base de planes", key="agregar_patologia_planes", use_container_width=True):
            clave_limpia = normalizar_texto(nueva_clave_patologia).upper().replace(" ", "_")
            label_limpio = nuevo_label_patologia.strip().upper()
            if not clave_limpia:
                st.error("Escribe un nombre corto o código para la patología.")
            elif not label_limpio:
                st.error("Escribe el nombre visible de la patología.")
            elif not nuevo_plan_patologia.strip():
                st.error("Escribe el plan de manejo para la nueva patología.")
            elif clave_limpia in planes_patologia:
                st.warning("Esa patología ya existe en la base de planes.")
            else:
                planes_patologia[clave_limpia] = nuevo_plan_patologia
                labels_planes_patologia[clave_limpia] = label_limpio
                guardar_planes_patologia(planes_patologia, labels_planes_patologia)
                st.success(f"Patología agregada: {label_limpio}")
                st.rerun()

        st.divider()
        st.markdown("**Editar patología existente**")
        clave_patologia_editor = st.selectbox(
            "Patología",
            list(labels_planes_patologia.keys()),
            format_func=lambda clave: labels_planes_patologia.get(clave, clave),
            key="plan_patologia_selector"
        )

        texto_plan_patologia = st.text_area(
            "Plan editable para esta patología",
            value=planes_patologia.get(clave_patologia_editor, ""),
            height=220,
            key=f"plan_patologia_editor_{clave_patologia_editor}"
        )

        col_plan_1, col_plan_2 = st.columns(2)
        if col_plan_1.button("Guardar plan de esta patología", key="guardar_plan_patologia", use_container_width=True):
            planes_patologia[clave_patologia_editor] = texto_plan_patologia
            guardar_planes_patologia(planes_patologia, labels_planes_patologia)
            st.success("Plan guardado correctamente.")

        if col_plan_2.button("Restablecer plan por defecto", key="restablecer_plan_patologia", use_container_width=True):
            planes_patologia[clave_patologia_editor] = PLANES_PATOLOGIA_DEFAULTS.get(clave_patologia_editor, "")
            guardar_planes_patologia(planes_patologia, labels_planes_patologia)
            st.session_state.pop(f"plan_patologia_editor_{clave_patologia_editor}", None)
            st.rerun()

    with st.expander("Base de datos de dosis de medicación", expanded=False):
        dosis_medicacion = cargar_dosis_medicacion()

        st.markdown("**Agregar medicamento nuevo**")
        col_nuevo_1, col_nuevo_2 = st.columns(2)
        with col_nuevo_1:
            nuevo_medicamento = st.text_input(
                "Nombre del medicamento",
                key="nuevo_medicamento_nombre"
            )
            nuevo_mg_kg = st.number_input(
                "Nuevo mg/kg/dosis",
                min_value=0.0,
                step=0.1,
                key="nuevo_medicamento_mgkg"
            )
            nuevo_intervalo = st.number_input(
                "Nuevo intervalo (horas)",
                min_value=1,
                step=1,
                key="nuevo_medicamento_intervalo"
            )
        with col_nuevo_2:
            nuevo_max_mg = st.number_input(
                "Nueva dosis máxima (mg)",
                min_value=0.0,
                step=1.0,
                key="nuevo_medicamento_maxmg"
            )
            nuevo_via = st.text_input(
                "Nueva vía",
                key="nuevo_medicamento_via"
            )
            nueva_indicacion = st.text_input(
                "Nueva indicación",
                key="nuevo_medicamento_indicacion"
            )

        if st.button("Agregar medicamento a la base de datos", key="agregar_medicamento_bd", use_container_width=True):
            nombre_normalizado = normalizar_texto(nuevo_medicamento).upper().replace(" ", "_")
            if not nuevo_medicamento.strip():
                st.error("Escribe un nombre para el medicamento.")
            elif nombre_normalizado in dosis_medicacion:
                st.warning("Ese medicamento ya existe en la base de datos.")
            else:
                dosis_medicacion[nombre_normalizado] = {
                    "mg_kg_dosis": float(nuevo_mg_kg),
                    "intervalo_horas": int(nuevo_intervalo),
                    "max_mg": float(nuevo_max_mg),
                    "via": nuevo_via.strip().upper(),
                    "indicacion": nueva_indicacion.strip().upper(),
                }
                guardar_dosis_medicacion(dosis_medicacion)
                st.success(f"Medicamento agregado: {nombre_normalizado}")
                st.rerun()

        st.divider()
        st.markdown("**Editar medicamento existente**")
        medicamento_editor = st.selectbox(
            "Medicamento",
            list(dosis_medicacion.keys()),
            key="medicacion_dosis_selector"
        )

        configuracion_actual = dosis_medicacion[medicamento_editor]
        col_dosis_1, col_dosis_2 = st.columns(2)
        with col_dosis_1:
            mg_kg_dosis = st.number_input(
                "mg/kg/dosis",
                min_value=0.0,
                value=float(configuracion_actual.get("mg_kg_dosis", 0.0)),
                step=0.1,
                key=f"mgkg_{medicamento_editor}"
            )
            intervalo_horas = st.number_input(
                "Intervalo (horas)",
                min_value=1,
                value=int(configuracion_actual.get("intervalo_horas", 6)),
                step=1,
                key=f"intervalo_{medicamento_editor}"
            )
        with col_dosis_2:
            max_mg = st.number_input(
                "Dosis máxima (mg)",
                min_value=0.0,
                value=float(configuracion_actual.get("max_mg", 0.0)),
                step=1.0,
                key=f"maxmg_{medicamento_editor}"
            )
            via = st.text_input(
                "Vía",
                value=str(configuracion_actual.get("via", "")),
                key=f"via_{medicamento_editor}"
            )

        indicacion = st.text_input(
            "Indicación",
            value=str(configuracion_actual.get("indicacion", "")),
            key=f"indicacion_{medicamento_editor}"
        )

        col_med_1, col_med_2 = st.columns(2)
        if col_med_1.button("Guardar dosis de este medicamento", key="guardar_dosis_medicacion", use_container_width=True):
            dosis_medicacion[medicamento_editor] = {
                "mg_kg_dosis": mg_kg_dosis,
                "intervalo_horas": intervalo_horas,
                "max_mg": max_mg,
                "via": via.strip().upper(),
                "indicacion": indicacion.strip().upper(),
            }
            guardar_dosis_medicacion(dosis_medicacion)
            st.success("Dosis guardada correctamente.")

        if col_med_2.button("Restablecer dosis por defecto", key="restablecer_dosis_medicacion", use_container_width=True):
            if medicamento_editor in DOSIS_MEDICACION_DEFAULTS:
                dosis_medicacion[medicamento_editor] = json.loads(json.dumps(DOSIS_MEDICACION_DEFAULTS[medicamento_editor]))
                guardar_dosis_medicacion(dosis_medicacion)
                st.rerun()
            else:
                st.info("Este medicamento no tiene un valor por defecto predefinido.")
